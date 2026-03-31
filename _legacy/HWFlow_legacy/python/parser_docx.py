"""
parser_docx.py — .docx 파일을 중간표현(IR)으로 변환

python-docx를 이용하여 문서 구조와 스타일 정보를 추출한다.

기능:
1. 문서 내용을 IR 블록으로 변환
2. 각 단락의 원본 스타일 정보를 JSON으로 반환 (인스펙터용)

Word 스타일 → IR 매핑:
  Heading 1  → heading1
  Heading 2  → heading2
  Heading 3  → heading3
  Heading 4  → heading4
  Normal     → body
  Table      → table
"""

from typing import List, Dict, Any, Optional, Tuple
from docx import Document
from docx.shared import Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


# Word 스타일 이름 → IR 타입 매핑
STYLE_MAP = {
    "Heading 1": "heading1",
    "Heading 2": "heading2",
    "Heading 3": "heading3",
    "Heading 4": "heading4",
    "Title": "heading1",
    "Subtitle": "heading2",
    "Normal": "body",
    "Body Text": "body",
    "List Paragraph": "body",
    # 한글 스타일명
    "제목 1": "heading1",
    "제목 2": "heading2",
    "제목 3": "heading3",
    "제목 4": "heading4",
    "본문": "body",
}


def parse_docx(file_path: str) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
    """
    .docx 파일을 파싱하여 (IR 블록 리스트, 스타일 정보 리스트) 튜플을 반환한다.

    Returns:
        (ir_blocks, style_infos)
        - ir_blocks: 변환용 중간표현
        - style_infos: 인스펙터 표시용 원본 스타일 정보
    """
    doc = Document(file_path)
    ir_blocks: List[Dict[str, Any]] = []
    style_infos: List[Dict[str, Any]] = []

    for element in doc.element.body:
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag

        if tag == 'p':
            # 일반 단락
            para = _find_paragraph(doc, element)
            if para is None:
                continue

            ir_type = _map_style(para.style.name if para.style else "Normal")
            runs = _extract_runs(para)

            # 빈 단락도 포함 (줄바꿈 용도)
            ir_blocks.append({"type": ir_type, "runs": runs})
            style_infos.append(_extract_style_info(para))

        elif tag == 'tbl':
            table = _find_table(doc, element)
            if table is None:
                continue

            table_block = _parse_table(table)
            ir_blocks.append(table_block)
            style_infos.append({"type": "table", "rows": len(table.rows), "cols": len(table.columns)})

    return ir_blocks, style_infos


def _find_paragraph(doc: Document, element) -> Optional[Any]:
    """element에 대응하는 Paragraph 객체를 찾는다."""
    for para in doc.paragraphs:
        if para._element is element:
            return para
    return None


def _find_table(doc: Document, element) -> Optional[Any]:
    """element에 대응하는 Table 객체를 찾는다."""
    for table in doc.tables:
        if table._element is element:
            return table
    return None


def _map_style(style_name: str) -> str:
    """Word 스타일 이름을 IR 타입으로 변환한다."""
    return STYLE_MAP.get(style_name, "body")


def _extract_runs(para) -> List[Dict[str, Any]]:
    """단락의 run들을 IR run 리스트로 변환한다."""
    runs = []
    for run in para.runs:
        text = run.text
        if not text:
            continue
        runs.append({
            "text": text,
            "bold": bool(run.bold),
            "italic": bool(run.italic) if run.italic else False,
            "underline": bool(run.underline) if run.underline else False,
            "color": _color_to_hex(run.font.color.rgb) if run.font.color and run.font.color.rgb else None,
        })

    if not runs and para.text:
        runs.append({"text": para.text, "bold": False})

    return runs


def _extract_style_info(para) -> Dict[str, Any]:
    """단락의 원본 스타일 정보를 추출한다 (인스펙터용)."""
    style_info: Dict[str, Any] = {
        "style_name": para.style.name if para.style else "Normal",
        "ir_mapping": _map_style(para.style.name if para.style else "Normal"),
        "text_preview": para.text[:50] if para.text else "",
    }

    # 단락 서식
    pf = para.paragraph_format
    if pf:
        align_map = {
            WD_ALIGN_PARAGRAPH.LEFT: "left",
            WD_ALIGN_PARAGRAPH.CENTER: "center",
            WD_ALIGN_PARAGRAPH.RIGHT: "right",
            WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
        }
        style_info["paragraph"] = {
            "alignment": align_map.get(pf.alignment, "unknown") if pf.alignment is not None else None,
            "line_spacing": float(pf.line_spacing) if pf.line_spacing else None,
            "space_before_pt": float(pf.space_before.pt) if pf.space_before else None,
            "space_after_pt": float(pf.space_after.pt) if pf.space_after else None,
            "first_line_indent_pt": float(pf.first_line_indent.pt) if pf.first_line_indent else None,
            "left_indent_pt": float(pf.left_indent.pt) if pf.left_indent else None,
        }

    # 첫 번째 run의 글자 서식
    if para.runs:
        run = para.runs[0]
        font = run.font
        style_info["font"] = {
            "name": font.name,
            "size_pt": float(font.size.pt) if font.size else None,
            "bold": font.bold,
            "italic": font.italic,
            "underline": font.underline is not None and font.underline is not False,
            "color": _color_to_hex(font.color.rgb) if font.color and font.color.rgb else None,
        }

    return style_info


def _parse_table(table) -> Dict[str, Any]:
    """Word 표를 IR table 블록으로 변환한다."""
    rows = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            runs = []
            for para in cell.paragraphs:
                if runs:
                    # 단락 사이 줄바꿈
                    runs.append({"text": "\n", "bold": False})
                if para.runs:
                    for r in para.runs:
                        if not r.text:
                            continue
                        runs.append({
                            "text": r.text,
                            "bold": bool(r.bold),
                        })
                elif para.text:
                    runs.append({"text": para.text, "bold": False})
            if not runs:
                runs.append({"text": "", "bold": False})
            cells.append({"runs": runs})
        rows.append(cells)

    # 첫 행이 헤더인지 추정 (첫 행의 모든 run이 bold이면 헤더)
    has_header = False
    if rows:
        first_row = rows[0]
        if all(
            all(r.get("bold", False) for r in cell.get("runs", []) if r.get("text", "").strip())
            for cell in first_row
            if any(r.get("text", "").strip() for r in cell.get("runs", []))
        ):
            has_header = True

    return {"type": "table", "rows": rows, "has_header": has_header}


def _color_to_hex(rgb) -> Optional[str]:
    """RGBColor를 #RRGGBB 문자열로 변환한다."""
    if rgb is None:
        return None
    return f"#{rgb}"


def get_style_report(file_path: str) -> List[Dict[str, Any]]:
    """
    인스펙터용: .docx 파일의 모든 단락에 대한 스타일 정보만 반환한다.
    """
    _, style_infos = parse_docx(file_path)
    return style_infos
